"""§4.7 Export — local file generation only. No upload, no Google Docs
API, no additional OAuth scope; python-docx runs entirely offline.

Content is built once into a format-agnostic ExportContent, then passed
through per-format adapter functions (.md, .txt, .docx) — adding a format
later means adding one adapter, not touching how content is assembled,
and every format stays consistent with the others by construction.

Covers both AutobioEntry and AutobioSummary (see content_for_autobio_entry
/ content_for_autobio_summary), per the spec's own requirement that export
apply to both.

Not built: the spec's optional direct-to-Obsidian-vault write for .md,
described as "consistent with the existing timer note skill's append
pattern in Nexus" — Nexus and that skill are outside this repo, with no
visible append-pattern convention to match here. The spec's own fallback
applies instead: ".md ... [o]therwise just saved as a file the user can
move themselves" — which is exactly what this does, for every format.
"""

from __future__ import annotations

import io
from dataclasses import dataclass


@dataclass
class ExportContent:
    title: str
    date_label: str
    body_text: str


def content_for_autobio_entry(entry: dict) -> ExportContent:
    """`entry` is whatever repository.get_autobio_entry() returns."""
    return ExportContent(
        title=f"Autobio — {entry['date']}",
        date_label=entry["date"],
        body_text=entry["final_text"],
    )


def content_for_autobio_summary(summary: dict) -> ExportContent:
    """`summary` is whatever repository.get_autobio_summary() returns."""
    return ExportContent(
        title=f"Autobio — {summary['start_date']} to {summary['end_date']}",
        date_label=f"{summary['start_date']} – {summary['end_date']}",
        body_text=summary["text"],
    )


def markdown_adapter(content: ExportContent) -> str:
    return f"# {content.title}\n\n*{content.date_label}*\n\n{content.body_text}\n"


def text_adapter(content: ExportContent) -> str:
    return f"{content.title}\n{content.date_label}\n\n{content.body_text}\n"


def docx_adapter(content: ExportContent) -> bytes:
    from docx import Document  # imported lazily: only this adapter needs it

    doc = Document()
    doc.add_heading(content.title, level=1)
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(content.date_label)
    date_run.italic = True

    for paragraph in content.body_text.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# (mimetype, adapter) — the mimetype travels with the adapter since a
# caller (the web route) needs both and they're intrinsically paired.
ADAPTERS = {
    "md": ("text/markdown", markdown_adapter),
    "txt": ("text/plain", text_adapter),
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        docx_adapter,
    ),
}


def export(content: ExportContent, fmt: str) -> tuple[str, bytes]:
    """Returns (mimetype, file bytes). Raises ValueError for an unknown
    format rather than silently picking one — a caller mistyping a format
    string should see that immediately, not get a wrong file."""
    if fmt not in ADAPTERS:
        raise ValueError(f"Unknown export format: {fmt!r} (expected one of {sorted(ADAPTERS)})")
    mimetype, adapter = ADAPTERS[fmt]
    result = adapter(content)
    data = result if isinstance(result, bytes) else result.encode("utf-8")
    return mimetype, data
